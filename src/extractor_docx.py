from __future__ import annotations

import sys
from pathlib import Path

from models import ExtractionResult


def extract_docx(path: Path, max_chars: int = 8000) -> ExtractionResult:
    try:
        from docx import Document

        document = Document(path)
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
        table_lines: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    table_lines.append(" | ".join(cells))
        combined = "\n".join(paragraphs + table_lines).strip()
        if not combined:
            return ExtractionResult(file_type="docx", extraction_status="empty_text", notes=["No text extracted from DOCX."])
        return ExtractionResult(
            file_type="docx",
            extraction_status="success",
            extracted_text=combined,
            text_excerpt=combined[:max_chars],
            notes=["DOCX text extracted."],
        )
    except ModuleNotFoundError:
        return ExtractionResult(file_type="docx", extraction_status="missing_dependency:python-docx", notes=["python-docx is not installed."])
    except Exception as exc:
        if sys.platform != "win32":
            return ExtractionResult(file_type="docx", extraction_status=f"error:{exc}", notes=["DOCX extraction failed."])
        # Windows: MIP 보호 파일일 수 있으므로 Word COM으로 재시도
        return _extract_docx_via_com(path, max_chars, fallback_error=str(exc))


from com_lock import word_lock

def _extract_docx_via_com(path: Path, max_chars: int, fallback_error: str) -> ExtractionResult:
    """MIP-보호 DOCX를 Word COM으로 열어 텍스트 추출."""
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return ExtractionResult(
            file_type="docx",
            extraction_status=f"error:{fallback_error}",
            notes=["DOCX extraction failed. pywin32 not installed for COM fallback."],
        )

    word = None
    com_initialized = False
    try:
        try:
            import pythoncom
            pythoncom.CoInitialize()
            com_initialized = True
        except Exception:
            pass
            
        with word_lock:
            word = win32com.client.Dispatch("Word.Application")
            try:
                word.Visible = False
            except Exception:
                pass
            try:
                word.DisplayAlerts = 0
            except Exception:
                pass
            
            doc = word.Documents.Open(str(path.resolve()), ReadOnly=True)
            # Range().Text는 단락, 표 셀, 머리글/바닥글 포함 전체 텍스트 반환
            # Word는 줄바꿈에 \r을 사용하므로 \n으로 정규화
            raw = doc.Range().Text
            doc.Close(False)
            
        combined = raw.replace("\r", "\n").strip()
        if not combined:
            return ExtractionResult(
                file_type="docx",
                extraction_status="empty_text",
                notes=["COM fallback: No text extracted from DOCX."],
            )
        return ExtractionResult(
            file_type="docx",
            extraction_status="success",
            extracted_text=combined,
            text_excerpt=combined[:max_chars],
            notes=["COM fallback (MIP): DOCX text extracted."],
        )
    except Exception as com_exc:
        return ExtractionResult(
            file_type="docx",
            extraction_status=f"error:{fallback_error}",
            notes=[f"DOCX extraction failed. COM fallback also failed: {com_exc}"],
        )
    finally:
        if word is not None:
            try:
                with word_lock:
                    word.Quit()
            except Exception:
                pass
        if com_initialized:
            try:
                import pythoncom
                pythoncom.CoUninitialize()
            except Exception:
                pass
